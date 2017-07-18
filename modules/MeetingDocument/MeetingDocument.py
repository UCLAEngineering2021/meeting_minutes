#imports
from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os.path
from pathlib import Path
from meeting_minutes.modules.input.MeetingInputParser import MeetingInputParser

#MeetingDocument

#Instance Data
###############
#docPath: string stores path to the document.  Defined relative to the placement of MeetingDocument.py for portability
#document: python docx object storing all document info
#timeCreated: string stores the timestamp marking when the meeting document was created
#author: string name of the author 'creating' the meeting document
#proceedings: string describes the events of the meeting
#location: string describes the location of the meeting
#date: string formatted like so: MM/DD/YY
#startTime: string stores the meeting's start time HH:MM
#endTime: string stores the time when the meeting ended HH:MM
#finances: string/string array stores potential financial transactions between last meeting and now
#students: string array stores names of students that attended meeting

#Static Class Data
TITLE = 'Miramonte EECS Club Meeting Minutes'
LINE_BREAKS_AFTER_TITLE = 2
FILES_DIR = 'files'
FILE_STORAGE_PATH = Path(__file__).parents[2]/FILES_DIR
DEFAULT_MEETING_NOTES_NAME = 'EECS_Club_Meeting_Notes_'
DOCX_EXT = '.docx'
HOME_DIR = '~'
RES_FILE_PATH = '/Desktop/meeting_minutes/res/'
LOGO_NAME = 'logo'
LOGO_EXT = '.png'

def _isSetToValue(param):
    if(param is not None):
        return True
    else:
        return False

class MeetingDocument:

    #Methods

    #constructor_method
    def __init__(self, argumentParser):
        #initialize list of all member variables for validity checking purposes in readyToWrite()
        self.varDict = []
        # get the author
        print(argumentParser.getUploaderName())
        self.author = argumentParser.getUploaderName()
        self.varDict.append(self.author)
        # get the date
        print(argumentParser.getDate())
        self.date = argumentParser.getDate()
        self.varDict.append(self.date)
        # get the beginning time
        print(argumentParser.getStartTime())
        self.startTime = argumentParser.getStartTime()
        self.varDict.append(self.startTime)
        # get the ending time
        print(argumentParser.getEndTime())
        self.endTime = argumentParser.getEndTime()
        self.varDict.append(self.endTime)
        # get the location
        print(argumentParser.getLocation())
        self.location = argumentParser.getLocation()
        self.varDict.append(self.location)
        # get the students present
        print(argumentParser.getStudents())
        self.students = argumentParser.getStudents()
        self.varDict.append(self.students)
        # get the students absent to the meeting
        print(argumentParser.getAbsentStudents())
        self.absentStudents = argumentParser.getAbsentStudents()
        self.varDict.append(self.absentStudents)
        # get the proceedings
        print(argumentParser.getProceedings())
        self.proceedings = argumentParser.getProceedings()
        self.varDict.append(self.proceedings)
        # get financial messages, if any
        print(argumentParser.getFinances())
        self.finances = argumentParser.getFinances()
        self.varDict.append(self.finances)
        # retrieve the current timestamp to append to this document's new filename
        timestamp = str(datetime.datetime.now().strftime('%Y-%m-%d_%H:%M:%S'))
        # get the current time, and store this into timeCreated
        self.timeCreated = timestamp
        # make the name
        self.docName = DEFAULT_MEETING_NOTES_NAME + timestamp + DOCX_EXT
        # initialize + open the document
        self.document = Document()
        # save the document
        self.docPath =  FILE_STORAGE_PATH/self.docName
        self.document.save(str(self.docPath))

    #returns true if all necessary fields have been supplied to this document by the argument parser
    #returns false otherwise
    def readyToWrite(self):
        # grab the list 'varDict' of all Meeting Document's member variables
        # loop through and make sure that each is properly set to a value
        for variableValue in self.varDict:
            if(variableValue == None):
                return False
        return True

    def _write(self, msg, centering=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        # add a paragraph with msg to the document
        p = self.document.add_paragraph()
        if(bold == True):
            # make the message bold, if necessary
            p.add_run(msg).bold = True
        else:
            p.add_run(msg)
        # adjust the centering
        p.alignment = centering
        self.document.save(str(self.docPath))

    def _writeLineBreak(self):
        # write a message just containing a newline character, with no special formatting
        self._write('\n')

    def _writeAuthor(self):
        # define author message
        AUTHOR_MSG = 'Author: '
        # write it (author message, right alignment, no bold)
        self._write(AUTHOR_MSG + self.author, WD_ALIGN_PARAGRAPH.RIGHT)
        # write a line break
        #self._writeLineBreak()

    def _writeTitle(self):
        # write it (title message, center alignment, bold)
        self._write(TITLE + ', ' + self.date, WD_ALIGN_PARAGRAPH.CENTER, True)
        # write a line break x2
        #for i in range(0, LINE_BREAKS_AFTER_TITLE):
            #self._writeLineBreak()

    def _writeSectionHeader(self, headerMessage):
        # write header message (message,bold, underline)
        self._write(headerMessage + ':')
        # write a line break
        #self._writeLineBreak()

    def _writeSection(self, headerMessage, sectionMessage):
        self._writeSectionHeader(headerMessage)
        self._write(sectionMessage)
        #self._writeLineBreak()

    def _writeBulletList(self, stringArray):
        bullet_paragraph = self.document.add_paragraph()
        for i in range(0, len(stringArray)):
            # append a newline onto every element
            arrayItem = stringArray[i] + '\n'
            # write a new paragraph with the style 'ListBullet' (Bullet List)
            bullet_paragraph.add_run(arrayItem)

    def _writeBulletSection(self, headerMessage, sectionList):
        self._writeSectionHeader(headerMessage)
        self._writeBulletList(sectionList)

    def _addPicture(self):
        homeDir = os.path.expanduser(HOME_DIR)
        picturePath = homeDir + RES_FILE_PATH + LOGO_NAME + LOGO_EXT
        if(os.path.isfile(picturePath)):
            print('picture has path')
        self.document.add_picture(picturePath)

    def populate(self):
        # To create a well formatted documentm writeAuthor and writeTitle must be executed in that order,
        # followed by any ordering of writeSection
        if(self.readyToWrite()):
            self._writeAuthor()
            self._addPicture()
            self._writeTitle()
            self._writeSection('Meeting Start Time', self.startTime)
            self._writeSection('Meeting End Time', self.endTime)
            self._writeSection('Meeting Location', self.location)
            self._writeBulletSection('Students Present', self.students)
            self._writeBulletSection('Students Absent', self.absentStudents)
            self._writeSection('Meeting Proceedings', self.proceedings)
            self._writeSection("Treasurer's Report", self.finances)
        else:
            print('Refusing to populate document!  Some Fields are unset!')

    def getTimeCreated(self):
        return self.timeCreated

    def getAuthor(self):
        return self.author

    def getName(self):
        return self.docName

    def getPath(self):
        return self.docPath
